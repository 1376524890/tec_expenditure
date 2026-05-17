#!/usr/bin/env python3
"""
Generate the course paper Word document:
研发费用加计扣除政策、研发投入调整与企业创新效率
——基于2021年制造业差异化税收激励的实证分析
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

OUTPUT_DIR = 'outputs'
FIGURE_DIR = os.path.join(OUTPUT_DIR, 'figures')

# ============================================================
# Helper functions
# ============================================================

def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False):
    """Set font for a run."""
    run.font.size = size
    run.bold = bold
    run.font.name = en_font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), cn_font)


def add_paragraph_with_font(doc, text, cn_font='宋体', en_font='Times New Roman',
                             size=Pt(12), bold=False, alignment=None,
                             first_line_indent=None, space_after=Pt(0), space_before=Pt(0),
                             line_spacing=1.5):
    """Add a paragraph with specific font settings."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    return p


def add_heading_custom(doc, text, level=1):
    """Add a custom heading. level 0=title, 1=一级, 2=二级, 3=三级"""
    if level == 0:
        # Main title - 黑体 小二加粗居中
        p = add_paragraph_with_font(doc, text, cn_font='黑体', en_font='Times New Roman',
                                     size=Pt(18), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     line_spacing=1.5, space_after=Pt(6))
    elif level == 1:
        # 一级标题 - 黑体四号加粗
        p = add_paragraph_with_font(doc, text, cn_font='黑体', en_font='Times New Roman',
                                     size=Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     line_spacing=1.5, space_before=Pt(12), space_after=Pt(6))
    elif level == 2:
        # 二级标题 - 楷体加粗小四号
        p = add_paragraph_with_font(doc, text, cn_font='楷体', en_font='Times New Roman',
                                     size=Pt(12), bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     line_spacing=1.5, space_before=Pt(6), space_after=Pt(3))
    elif level == 3:
        # 三级标题 - 宋体加粗小四号
        p = add_paragraph_with_font(doc, text, cn_font='宋体', en_font='Times New Roman',
                                     size=Pt(12), bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                     line_spacing=1.5, space_before=Pt(6), space_after=Pt(3))
    return p


def add_body_para(doc, text, first_indent=True):
    """Add body text paragraph. 宋体小四号."""
    indent = Cm(0.74) if first_indent else None
    return add_paragraph_with_font(doc, text, cn_font='宋体', en_font='Times New Roman',
                                    size=Pt(12), bold=False,
                                    first_line_indent=indent,
                                    line_spacing=1.5, space_after=Pt(0))


def add_table_title(doc, text):
    """Add table title - centered, 宋体小五号加粗."""
    return add_paragraph_with_font(doc, text, cn_font='宋体', en_font='Times New Roman',
                                    size=Pt(9), bold=True,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    line_spacing=1.5, space_before=Pt(6), space_after=Pt(3))


def add_figure_title(doc, text):
    """Add figure title below figure - centered, 宋体小五号加粗."""
    return add_paragraph_with_font(doc, text, cn_font='宋体', en_font='Times New Roman',
                                    size=Pt(9), bold=True,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    line_spacing=1.5, space_before=Pt(3), space_after=Pt(6))


def add_table_note(doc, text):
    """Add table note - 宋体小五号."""
    return add_paragraph_with_font(doc, text, cn_font='宋体', en_font='Times New Roman',
                                    size=Pt(9), bold=False,
                                    line_spacing=1.25, space_after=Pt(6))


def create_three_line_table(doc, headers, rows, col_widths=None):
    """Create a three-line table (三线表). Returns the table object."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5), bold=True)

    # Set data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = ''
            p = row_cells[c_idx].paragraphs[0]
            # Right-align numbers, left-align text
            if isinstance(val, (int, float)):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5))

    # Apply three-line border style
    # Remove all borders first, then add top, bottom, and header-bottom lines
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)

    # Top border (thick) on first row
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": 12, "color": "000000"})

    # Bottom border on header row (medium)
    for cell in table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": 6, "color": "000000"})

    # Bottom border (thick) on last row
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": 12, "color": "000000"})

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ============================================================
# Main document generation
# ============================================================

def generate_paper():
    doc = Document()

    # ---- Page setup ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ---- Default style ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ============================================================
    # TITLE
    # ============================================================
    add_paragraph_with_font(doc, '', size=Pt(12), line_spacing=1.5)  # spacer

    add_heading_custom(doc, '研发费用加计扣除政策、研发投入调整与企业创新效率', level=0)

    add_paragraph_with_font(
        doc, '——基于2021年制造业差异化税收激励的实证分析',
        cn_font='楷体', en_font='Times New Roman', size=Pt(12), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=Pt(12)
    )

    add_paragraph_with_font(
        doc, '作者：XXX',
        cn_font='仿宋', en_font='Times New Roman', size=Pt(14), bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=Pt(18)
    )

    # ============================================================
    # ABSTRACT
    # ============================================================
    add_heading_custom(doc, '摘要', level=1)

    abstract_text = (
        '研发费用加计扣除是我国支持企业创新的重要税收激励工具。基于2017至2022年A股上市公司面板数据，'
        '利用2021年制造业企业研发费用加计扣除比例由75%提高至100%的政策变化，采用双重差分方法考察税收激励政策'
        '是否伴随企业研发投入调整和相对创新效率改善。研究发现，政策后制造业企业研发投入和专利产出在绝对水平上'
        '仍保持增长，但相对于非制造业企业，制造业专利数量变化并不显著，研发支出、研发强度和研发人员投入增长'
        '显著较慢。在创新产出未显著扩张、研发投入相对放缓的背景下，制造业企业单位研发投入和单位研发人员创新'
        '产出效率表现出显著相对改善。效率来源拆解表明，创新效率改善主要来自研发投入相对增长较慢，而非专利数量'
        '显著增加。异质性分析显示，政策前研发基础较强的企业和制造业内部高研发暴露企业效率改善更明显。稳健性'
        '检验表明，加入省份与年份交互固定效应后核心结果保持稳定；但事件研究和安慰剂检验提示研发投入和效率指标'
        '存在政策前趋势差异，因此本文将结论限定为政策伴随研发投入调整和相对创新效率改善，不作严格因果意义上的'
        '效率提升判断。研究表明，税收激励政策的评估不宜仅关注专利数量扩张，还应纳入研发投入结构、单位投入产出'
        '效率和政策识别边界。'
    )
    add_body_para(doc, abstract_text, first_indent=True)

    # Keywords
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(0.74)
    run_label = p.add_run('关键词：')
    set_run_font(run_label, cn_font='黑体', en_font='Times New Roman', size=Pt(12), bold=True)
    kw_text = '研发费用加计扣除；税收激励；研发投入；创新效率；双重差分'
    run_kw = p.add_run(kw_text)
    set_run_font(run_kw, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False)

    # ============================================================
    # 一、数据来源与变量设定
    # ============================================================
    add_heading_custom(doc, '一、数据来源与变量设定', level=1)

    add_body_para(doc,
        '本文使用2017至2022年A股上市公司面板数据，样本对象为沪深A股上市公司。企业层面数据主要来自'
        '国泰安（CSMAR）数据库，涉及利润表、资产负债表、现金流量表、上市公司基本信息年度表、国内外专利'
        '申请获得情况表、研发投入情况表、政府补助表和上市公司控制人文件。财务数据限定为合并报表口径，'
        '仅保留12月31日年报数据。宏观政策和研发投入背景数据来自国家统计局、财政部、税务总局等官方渠道。'
    )

    add_body_para(doc,
        '最终样本包括5404家企业、26772条企业年度观测值，其中制造业观测值为16740条，占比62.5%。'
        '回归有效样本因变量口径略有差异，以发明专利申请模型为例，最终有效观测值为25831条。制造业识别'
        '依据为证监会2012年行业分类标准，行业代码以C开头的企业定义为制造业企业。'
    )

    add_body_para(doc,
        '专利数据来自CSMAR国内外专利申请获得情况表，主要包括发明专利申请、发明专利授权、专利总申请和'
        '专利总授权。研发投入数据来自CSMAR研发投入情况表，主要包括研发支出、研发强度、研发人员和研发人员'
        '占比。财务控制变量包括企业规模、盈利能力、现金流比率和企业年龄。数据处理方面，专利数量采用'
        'ln(1+x)处理，研发支出采用万元口径并进行ln(1+x)处理，研发人员采用ln(1+x)处理。连续变量按年度'
        '1%和99%分位数缩尾，以降低极端值对估计结果的影响。资产负债率变量不可用，本文使用现金流比率作为'
        '补充控制变量。'
    )

    # Table 1: Sample selection
    add_table_title(doc, '表 1 样本筛选流程表')
    table1 = create_three_line_table(doc,
        ['步骤', '处理内容', '删除观测值', '剩余观测值'],
        [
            ['1', '初始A股上市公司样本（2017至2022年合并后）', '—', '26,772'],
            ['2', '剔除ST及*ST企业', '当前文件无法追溯', '—'],
            ['3', '剔除金融行业', '当前文件无法追溯', '—'],
            ['4', '剔除核心变量缺失观测', '因模型而异', '—'],
            ['5', '连续变量按年度1%和99%缩尾', '0', '26,772'],
            ['6', '最终回归样本（以发明专利申请模型为例）', '941', '25,831'],
        ]
    )
    add_table_note(doc, '注：ST企业和金融行业剔除步骤在数据合并阶段完成，当前材料无法追溯每一步删除数量。')

    # Table 2: Variable definitions
    add_table_title(doc, '表 2 主要变量定义表')
    table2 = create_three_line_table(doc,
        ['变量类型', '变量名称', '变量符号', '变量定义'],
        [
            ['创新数量', '发明专利申请', 'ln_invention_apply', 'ln(1+发明专利申请数量)'],
            ['创新数量', '发明专利授权', 'ln_invention_grant', 'ln(1+发明专利授权数量)'],
            ['创新数量', '专利总申请', 'ln_patent_apply_total', 'ln(1+三类专利申请合计)'],
            ['创新数量', '专利总授权', 'ln_patent_grant_total', 'ln(1+三类专利授权合计)'],
            ['研发投入', '研发支出（万元对数）', 'ln_rd_expense_10k', 'ln(1+研发支出/10000)'],
            ['研发投入', '研发强度', 'rd_intensity_01', '研发支出/营业收入（0—1口径）'],
            ['研发投入', '研发人员', 'ln_rd_staff', 'ln(1+研发人员数量)'],
            ['研发投入', '研发人员占比', 'rd_staff_ratio_01', '研发人员/员工总数（0—1口径）'],
            ['创新效率', '发明申请/研发支出效率', 'eff_apply_rd_10k', 'ln_invention_apply − ln_rd_expense_10k'],
            ['创新效率', '发明授权/研发支出效率', 'eff_grant_rd_10k', 'ln_invention_grant − ln_rd_expense_10k'],
            ['创新效率', '发明申请/研发人员效率', 'eff_apply_staff', 'ln_invention_apply − ln_rd_staff'],
            ['创新效率', '发明授权/研发人员效率', 'eff_grant_staff', 'ln_invention_grant − ln_rd_staff'],
            ['核心解释变量', '制造业×政策后', 'manufacturing_post2021', 'manufacturing × post2021'],
            ['控制变量', '企业规模', 'ln_assets', 'ln(总资产)'],
            ['控制变量', '盈利能力', 'roa', '净利润/总资产'],
            ['控制变量', '现金流比率', 'cashflow_ratio', '经营现金流/总资产'],
            ['控制变量', '企业年龄', 'firm_age', '年份减成立年份'],
        ]
    )
    add_table_note(doc, '注：完整变量定义见附录附表A。资产负债率变量不可用，使用现金流比率作为补充控制变量。')

    # ============================================================
    # 二、描述性统计分析
    # ============================================================
    add_heading_custom(doc, '二、描述性统计分析', level=1)

    add_body_para(doc,
        '从全样本描述性统计看，样本企业创新产出分布呈现明显右偏特征。发明专利申请对数均值为0.396，'
        '发明专利授权对数均值为0.307，四分位数显示相当部分企业在部分年份没有发明专利申请或授权。研发投入'
        '变量同样存在较大差异，研发支出万元对数均值为7.500，标准差为3.405，说明上市公司之间研发投入规模'
        '差异较大。研发强度均值为0.054，研发人员对数均值为4.468，研发人员占比均值为0.170。创新效率指标'
        '均值为负，主要由于效率变量采用专利产出对数减研发投入对数或研发人员对数构造，在研发投入规模较大'
        '的企业中该指标通常较低。'
    )

    # Table 3: Descriptive statistics
    add_table_title(doc, '表 3 全样本描述性统计表')
    desc_data = [
        ['ln_invention_apply', '26,772', '0.396', '1.118', '0.000', '0.000', '9.089'],
        ['ln_invention_grant', '26,772', '0.307', '0.902', '0.000', '0.000', '8.093'],
        ['ln_patent_apply_total', '26,772', '0.590', '1.490', '0.000', '0.000', '9.737'],
        ['ln_patent_grant_total', '26,772', '0.537', '1.395', '0.000', '0.000', '9.503'],
        ['ln_rd_expense_10k', '26,772', '7.500', '3.405', '0.000', '8.464', '15.420'],
        ['rd_intensity_01', '22,829', '0.054', '0.059', '0.000', '0.039', '0.401'],
        ['ln_rd_staff', '26,772', '4.468', '2.389', '0.000', '5.081', '11.152'],
        ['rd_staff_ratio_01', '21,953', '0.170', '0.139', '0.004', '0.135', '0.720'],
        ['eff_apply_rd_10k', '26,772', '−7.104', '3.378', '−15.420', '−8.036', '5.075'],
        ['eff_grant_rd_10k', '26,772', '−7.193', '3.345', '−14.870', '−8.091', '4.344'],
        ['ln_assets', '26,772', '22.217', '1.532', '19.041', '22.018', '27.203'],
        ['roa', '26,771', '0.041', '0.088', '−0.351', '0.045', '0.246'],
        ['cashflow_ratio', '26,770', '0.048', '0.072', '−0.172', '0.047', '0.252'],
        ['firm_age', '25,831', '20.269', '5.986', '7.000', '20.000', '38.000'],
    ]
    table3 = create_three_line_table(doc,
        ['变量', '观测值', '均值', '标准差', '最小值', '中位数', '最大值'],
        desc_data
    )
    add_table_note(doc, '注：所有连续变量已按年度1%和99%分位数缩尾。效率指标均值为负，原因参见正文说明。')

    add_body_para(doc,
        '从制造业年度均值看，制造业企业研发投入和专利产出在政策后没有出现绝对下降。制造业研发支出均值'
        '由2017年的1.64亿元提高至2022年的3.16亿元，研发强度由4.49%上升至6.28%，发明专利申请均值由'
        '10.8件提高至15.5件。该事实说明，后文双重差分结果中的研发投入下降应理解为制造业相对于非制造业'
        '的增速放缓，而非制造业企业研发投入绝对减少。'
    )

    # Table 4: Manufacturing annual trends
    add_table_title(doc, '表 4 制造业年度均值趋势表')
    trend_data = [
        ['2017', '1.64', '4.49', '10.8', '3.7', '478', '2,351'],
        ['2018', '2.09', '4.91', '11.5', '3.9', '535', '2,402'],
        ['2019', '2.28', '5.26', '15.4', '5.7', '558', '2,532'],
        ['2020', '2.34', '5.44', '12.7', '5.8', '542', '2,867'],
        ['2021', '2.79', '5.71', '13.3', '5.5', '558', '3,173'],
        ['2022', '3.16', '6.28', '15.5', '6.5', '594', '3,415'],
    ]
    table4 = create_three_line_table(doc,
        ['年份', '研发支出（亿元）', '研发强度（%）', '发明申请（件）',
         '发明授权（件）', '研发人员（人）', '观测数'],
        trend_data
    )

    # Insert Figure 1
    trend_png = os.path.join(FIGURE_DIR, 'focus_trend_rd_expense.png')
    if os.path.exists(trend_png):
        add_figure_title(doc, '图 1 制造业与非制造业研发支出变动趋势图')
        doc.add_picture(trend_png, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_note(doc, '数据来源：CSMAR数据库，作者计算。竖线标注2021年政策冲击时点。')

    add_body_para(doc,
        '制造业与非制造业分组比较显示，制造业企业在专利产出、研发支出和研发人员方面均显著高于非制造业'
        '企业。制造业企业发明专利申请对数均值为0.512，非制造业为0.181；制造业研发支出万元对数均值为'
        '8.672，非制造业为5.545。与此同时，制造业企业效率指标均值低于非制造业，这与制造业研发投入规模'
        '更大有关，也提示单纯比较效率均值不能识别政策影响，需要进一步使用固定效应模型控制企业不随时间'
        '变化的特征。分组比较详细结果见附录附表B。'
    )

    # ============================================================
    # 三、模型建立
    # ============================================================
    add_heading_custom(doc, '三、模型建立', level=1)

    add_body_para(doc,
        '本文利用2021年制造业企业研发费用加计扣除比例提高形成的行业差异化政策冲击构造双重差分模型。'
        '2018年至2020年，符合条件企业研发费用加计扣除比例提高至75%；2021年起，制造业企业研发费用加计'
        '扣除比例进一步提高至100%；2023年起，符合条件企业研发费用加计扣除比例统一提高至100%，制造业'
        '相对非制造业的政策强度差异随之减弱。2018年75%政策、2021年制造业100%政策和2023年全行业100%政策'
        '均有官方文件依据，其中2021年第13号公告明确制造业企业自2021年1月1日起按100%加计扣除，形成本文'
        '识别所依赖的政策差异。'
    )

    add_body_para(doc,
        '本文将制造业企业定义为处理组，非制造业企业定义为对照组，2021年及以后年份定义为政策后时期。'
        '基准模型设定为：'
    )

    # Formula (1)
    add_paragraph_with_font(
        doc, 'Yᵢₜ = α + β(Manufacturingᵢ × Post2021ₜ) + γXᵢₜ + μᵢ + λₜ + εᵢₜ   （1）',
        cn_font='宋体', en_font='Times New Roman', size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=Pt(6)
    )

    add_body_para(doc,
        '其中，Yᵢₜ表示企业i在年份t的创新数量、研发投入或创新效率变量。Manufacturingᵢ为制造业企业虚拟'
        '变量，Post2021ₜ为2021年及以后年份虚拟变量，二者交互项Manufacturingᵢ×Post2021ₜ的系数β是本文'
        '关注的核心估计量。Xᵢₜ表示控制变量，包括企业规模、盈利能力、现金流比率和企业年龄。模型同时控制'
        '企业固定效应μᵢ和年份固定效应λₜ，标准误在企业层面聚类。'
    )

    add_body_para(doc,
        '创新效率变量采用对数差分方式构造，其基本逻辑为：'
    )

    add_paragraph_with_font(
        doc, 'Efficiencyᵢₜ = ln(Innovationᵢₜ) − ln(RDInputᵢₜ)   （2）',
        cn_font='宋体', en_font='Times New Roman', size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=Pt(6)
    )

    add_body_para(doc,
        '因此，在相同样本和相同模型设定下，效率DID系数可以拆解为创新产出DID系数与研发投入DID系数之差：'
    )

    add_paragraph_with_font(
        doc, 'DID(Efficiency) ≈ DID(Innovation Output) − DID(RD Input)   （3）',
        cn_font='宋体', en_font='Times New Roman', size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=Pt(6)
    )

    add_body_para(doc,
        '这一设定的含义是，如果政策后制造业企业专利产出相对于非制造业没有显著变化，而研发投入相对增长'
        '较慢，则单位研发投入或单位研发人员对应的创新产出效率可能表现为相对改善。该解释依赖研发投入和'
        '创新产出两个维度的共同判断，因此后续需要结合事件研究和安慰剂检验审慎界定结论边界。'
    )

    # ============================================================
    # 四、模型结果分析
    # ============================================================
    add_heading_custom(doc, '四、模型结果分析', level=1)

    # (一) Innovation quantity
    add_heading_custom(doc, '（一）创新数量效应', level=2)

    add_body_para(doc,
        '基准DID结果显示，2021年制造业研发费用加计扣除比例提高后，制造业企业相对于非制造业企业的专利'
        '数量变化并不显著。以发明专利申请为因变量时，DID系数为−0.025，p值为0.3146；以发明专利授权为'
        '因变量时，DID系数为−0.011，p值为0.6224；专利总申请和专利总授权的DID系数同样不显著。该结果说明，'
        '政策后制造业企业并未在短期内表现出相对于非制造业的专利数量扩张。'
    )

    # (二) R&D adjustment
    add_heading_custom(doc, '（二）研发投入调整效应', level=2)

    add_body_para(doc,
        '研发投入变量呈现出更加明显的变化。以研发支出万元对数为因变量时，DID系数为−0.288，并在1%水平上'
        '显著；以研发强度为因变量时，DID系数为−0.003，并在5%水平上显著；以研发人员对数为因变量时，DID'
        '系数为−0.197，并在1%水平上显著。研发人员占比的系数接近于零且不显著。结合描述性事实可以看出，'
        '制造业企业研发投入并未绝对下降，其相对变化主要表现为政策后相对于非制造业的增长速度放缓。'
    )

    # (三) Innovation efficiency
    add_heading_custom(doc, '（三）创新效率效应', level=2)

    add_body_para(doc,
        '创新效率结果显示，单位研发投入和单位研发人员对应的创新产出效率均表现出显著相对改善。以发明申请'
        '与研发支出效率为因变量时，DID系数为0.264，并在1%水平上显著；以发明授权与研发支出效率为因变量时，'
        'DID系数为0.277，并在1%水平上显著；以发明申请与研发人员效率、发明授权与研发人员效率为因变量时，'
        'DID系数分别为0.172和0.186，均达到1%显著性水平。该结果表明，在专利数量未显著增加的情况下，制造业'
        '企业单位研发投入和单位研发人员产出的专利数量表现出相对提升。'
    )

    # Table 5: Baseline DID results
    add_table_title(doc, '表 5 基准双重差分估计结果表')
    did_data = [
        ['创新数量', 'ln_invention_apply', '−0.025', '0.0246', '0.3146', '25,831'],
        ['创新数量', 'ln_invention_grant', '−0.011', '0.0230', '0.6224', '25,831'],
        ['创新数量', 'ln_patent_apply_total', '−0.002', '0.0339', '0.9591', '25,831'],
        ['创新数量', 'ln_patent_grant_total', '0.035', '0.0357', '0.3298', '25,831'],
        ['研发投入', 'ln_rd_expense_10k', '−0.288***', '0.0568', '0.0000', '25,831'],
        ['研发投入', 'rd_intensity_01', '−0.003**', '0.0013', '0.0245', '22,822'],
        ['研发投入', 'ln_rd_staff', '−0.197***', '0.0443', '0.0000', '25,831'],
        ['研发投入', 'rd_staff_ratio_01', '0.000', '0.0029', '0.9509', '21,949'],
        ['创新效率', 'eff_apply_rd_10k', '0.264***', '0.0616', '0.0000', '25,831'],
        ['创新效率', 'eff_grant_rd_10k', '0.277***', '0.0611', '0.0000', '25,831'],
        ['创新效率', 'eff_apply_staff', '0.172***', '0.0499', '0.0005', '25,831'],
        ['创新效率', 'eff_grant_staff', '0.186***', '0.0491', '0.0002', '25,831'],
    ]
    table5 = create_three_line_table(doc,
        ['结果类别', '因变量', 'DID系数', '标准误', 'p值', 'N'],
        did_data
    )
    add_table_note(doc,
        '注：所有模型控制企业固定效应和年份固定效应，标准误在企业层面聚类。控制变量包括企业规模、ROA、'
        '现金流比率和企业年龄。*** p<0.01，** p<0.05，* p<0.10。'
    )

    # Insert Figure 2: Forest plot
    forest_png = os.path.join(FIGURE_DIR, 'focus_main_forest.png')
    if os.path.exists(forest_png):
        add_figure_title(doc, '图 2 三类结果变量DID系数森林图')
        doc.add_picture(forest_png, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_note(doc, '注：横线表示95%置信区间。创新数量系数置信区间穿过零，研发投入系数显著为负，创新效率系数显著为正。')

    # (四) Efficiency decomposition
    add_heading_custom(doc, '（四）效率来源拆解', level=2)

    add_body_para(doc,
        '效率来源拆解进一步说明，创新效率改善主要来自研发投入端的相对调整。以发明申请与研发支出效率为例，'
        '发明专利申请DID系数为−0.0247，研发支出DID系数为−0.2884，两者相减得到0.2637，与发明申请/研发支出'
        '效率DID系数完全一致。该恒等关系说明，效率改善并不来自专利数量显著增加，而是在专利产出相对稳定的'
        '情况下，研发投入相对增长较慢所形成的单位投入产出改善。'
    )

    # Table 6: Decomposition
    add_table_title(doc, '表 6 创新效率来源拆解表')
    decomp_data = [
        ['DID(ln_invention_apply)', '−0.025', '0.0246', '0.3146', '创新产出相对变化'],
        ['DID(ln_rd_expense_10k)', '−0.288***', '0.0568', '0.0000', '研发投入相对变化'],
        ['DID(eff_apply_rd_10k)', '0.264***', '0.0616', '0.0000', '单位研发投入效率相对变化'],
        ['DID(inv) − DID(rd)', '0.264', '—', '—', '与效率DID一致'],
    ]
    table6 = create_three_line_table(doc,
        ['项目', 'DID系数', '标准误', 'p值', '说明'],
        decomp_data
    )

    # Insert Figure 3: Efficiency decomposition
    decomp_png = os.path.join(FIGURE_DIR, 'focus_efficiency_decomposition.png')
    if os.path.exists(decomp_png):
        add_figure_title(doc, '图 3 创新效率来源拆解图')
        doc.add_picture(decomp_png, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_note(doc, '注：柱状图展示DID系数，横线为95%置信区间。效率DID约等于创新产出DID减去研发投入DID。')

    # (五) Heterogeneity
    add_heading_custom(doc, '（五）异质性分析', level=2)

    add_body_para(doc,
        '异质性分析显示，政策前研发基础较强的企业效率改善更明显。在高研发基础企业交互项模型中，'
        'did_x_high_pre_rd在四个效率指标上均为正，并达到统计显著水平。其中，发明申请/研发支出效率的'
        '交互项系数为0.225，发明授权/研发支出效率的交互项系数为0.272，说明研发基础较强的企业更容易将'
        '税收激励和研发资源调整转化为单位投入产出改善。制造业内部研发暴露检验也得到类似结果，高研发基础'
        '制造业企业在2021年后的效率改善更明显。'
    )

    # Table 7: Heterogeneity
    add_table_title(doc, '表 7 异质性分析结果表')
    het_data = [
        ['高研发基础', 'eff_apply_rd_10k', 'did_x_high_pre_rd', '0.225***', '0.0509', '0.0000'],
        ['高研发基础', 'eff_grant_rd_10k', 'did_x_high_pre_rd', '0.272***', '0.0484', '0.0000'],
        ['高研发基础', 'eff_apply_staff', 'did_x_high_pre_rd', '0.137**', '0.0534', '0.0104'],
        ['高研发基础', 'eff_grant_staff', 'did_x_high_pre_rd', '0.183***', '0.0512', '0.0003'],
        ['制造业内部', 'eff_apply_rd_10k', 'high_pre_rd_post2021', '0.198***', '0.0503', '0.0001'],
        ['制造业内部', 'eff_grant_rd_10k', 'high_pre_rd_post2021', '0.244***', '0.0476', '0.0000'],
        ['制造业内部', 'eff_apply_staff', 'high_pre_rd_post2021', '0.124**', '0.0537', '0.0215'],
        ['制造业内部', 'eff_grant_staff', 'high_pre_rd_post2021', '0.170***', '0.0514', '0.0009'],
    ]
    table7 = create_three_line_table(doc,
        ['异质性维度', '因变量', '交互项或处理变量', '系数', '标准误', 'p值'],
        het_data
    )
    add_table_note(doc,
        '注：高研发基础定义为政策前（2017至2020年）企业平均研发强度高于中位数。制造业内部检验仅使用'
        '制造业子样本。*** p<0.01，** p<0.05，* p<0.10。'
    )

    add_body_para(doc,
        '所有制异质性结果需要谨慎处理。完整材料显示，非国有企业分组样本较小（524家企业，2131条观测），'
        '交互项模型中非国企交互项多数不显著；分组回归中，非国企发明授权/研发人员效率系数为0.434，并在5%'
        '水平上显著。由于非国企样本量较小，正文中仅简略说明该结果提示市场化主体可能存在更强调整弹性，'
        '不作为核心结论展开。所有制异质性完整结果见附录附表C。'
    )

    # Insert Figure 4: Heterogeneity forest
    het_forest_png = os.path.join(FIGURE_DIR, 'focus_heterogeneity_forest.png')
    if os.path.exists(het_forest_png):
        add_figure_title(doc, '图 4 异质性结果森林图')
        doc.add_picture(het_forest_png, width=Inches(5.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_note(doc, '注：展示eff_apply_rd_10k异质性结果。横线为95%置信区间。')

    # ============================================================
    # 五、稳健性检验
    # ============================================================
    add_heading_custom(doc, '五、稳健性检验', level=1)

    # (一) Prov×Year FE
    add_heading_custom(doc, '（一）加入省份与年份交互固定效应', level=2)

    add_body_para(doc,
        '考虑到不同地区在财政科技支出、产业结构和创新环境方面可能存在年度冲击差异，本文进一步加入省份与'
        '年份交互固定效应。结果显示，核心效率指标保持显著。发明申请/研发支出效率的稳健性系数为0.255'
        '（p=0.0001），发明授权/研发支出效率的稳健性系数为0.290（p<0.0001），发明申请/研发人员效率和发明'
        '授权/研发人员效率也保持显著。研发支出和研发人员投入的DID系数仍显著为负。这说明，核心结果并非完全'
        '由省份层面的年度冲击驱动。'
    )

    # (二) Alternative efficiency
    add_heading_custom(doc, '（二）替代效率指标', level=2)

    add_body_para(doc,
        '本文尝试使用asinh比率型效率指标作为替代口径，但结果并不理想。asinh_apply_per_rd系数为负且在5%'
        '水平上显著，其余比率型指标多数不显著。该结果可能源于比率型指标对分母接近零的观测高度敏感，即使'
        '采用asinh变换也难以完全缓解偏态分布问题。因此，本文将对数差分型效率指标作为主结果，并将比率型'
        '效率指标作为稳健性讨论而非核心证据。'
    )

    # (三) Event study
    add_heading_custom(doc, '（三）事件研究检验', level=2)

    add_body_para(doc,
        '事件研究显示，创新产出变量的平行趋势相对较好。以发明专利申请为因变量时，2017年至2019年事件研究'
        '系数分别为0.003、−0.032和−0.014，均不显著，说明制造业和非制造业在政策前的发明专利申请趋势没有'
        '明显系统差异。然而，研发支出变量在政策前存在显著趋势差异，2017年至2019年系数分别为0.687、0.396'
        '和0.296，均在1%水平上显著。效率指标在政策前显著为负，2021年至2022年收敛至零附近。该模式说明，'
        '效率DID结果部分反映制造业政策前研发投入优势收窄，而非完全由2021年政策冲击导致。'
    )

    # Table 8: Event study
    add_table_title(doc, '表 8 事件研究核心结果表')
    event_data = [
        ['ln_invention_apply', '2017', '0.003', '0.033', '0.9252', '不显著'],
        ['ln_invention_apply', '2018', '−0.032', '0.031', '0.3042', '不显著'],
        ['ln_invention_apply', '2019', '−0.014', '0.030', '0.6393', '不显著'],
        ['ln_invention_apply', '2021', '−0.033', '0.028', '0.2288', '不显著'],
        ['ln_invention_apply', '2022', '−0.035', '0.031', '0.2634', '不显著'],
        ['ln_rd_expense_10k', '2017', '0.687***', '0.103', '0.0000', '显著'],
        ['ln_rd_expense_10k', '2018', '0.396***', '0.077', '0.0000', '显著'],
        ['ln_rd_expense_10k', '2019', '0.296***', '0.057', '0.0000', '显著'],
        ['ln_rd_expense_10k', '2021', '0.010', '0.055', '0.8541', '不显著'],
        ['ln_rd_expense_10k', '2022', '−0.002', '0.066', '0.9742', '不显著'],
        ['eff_apply_rd_10k', '2017', '−0.684***', '0.107', '0.0000', '显著'],
        ['eff_apply_rd_10k', '2018', '−0.428***', '0.083', '0.0000', '显著'],
        ['eff_apply_rd_10k', '2019', '−0.310***', '0.064', '0.0000', '显著'],
        ['eff_apply_rd_10k', '2021', '−0.044', '0.061', '0.4754', '不显著'],
        ['eff_apply_rd_10k', '2022', '−0.033', '0.072', '0.6484', '不显著'],
    ]
    table8 = create_three_line_table(doc,
        ['因变量', '年份', '系数', '标准误', 'p值', '结论'],
        event_data
    )
    add_table_note(doc, '注：基准年为2020年。*** p<0.01。N=25,831。')

    # Insert Figure 5: Event study
    event_png = os.path.join(FIGURE_DIR, 'focus_event_study.png')
    if os.path.exists(event_png):
        add_figure_title(doc, '图 5 事件研究系数图')
        doc.add_picture(event_png, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_note(doc, '注：基准年为2020年，虚线右侧为政策后时期。柱状范围为95%置信区间。')

    # (四) Placebo
    add_heading_custom(doc, '（四）安慰剂检验', level=2)

    add_body_para(doc,
        '安慰剂检验进一步提示因果解释需要谨慎。以2019年或2020年作为假想政策时点时，效率指标的DID系数仍然'
        '显著为正。假想2019年政策时点下，发明申请/研发支出效率系数为0.290（p=0.0001）；假想2020年政策时点'
        '下为0.325（p<0.0001）。与此同时，发明专利申请的假想DID系数并不显著，研发支出的假想DID系数显著为负。'
        '该结果与事件研究一致，说明制造业和非制造业在研发投入和效率方面的差异在政策前已经出现收敛趋势。'
        '基于此，本文将核心结论表述为"政策伴随研发投入调整和相对创新效率改善"，不将其解释为严格因果意义上'
        '的政策效率提升。安慰剂检验完整结果见附录附表D。'
    )

    # ============================================================
    # 六、政策分析
    # ============================================================
    add_heading_custom(doc, '六、政策分析', level=1)

    # (一) Policy evolution
    add_heading_custom(doc, '（一）政策演进与制度特征', level=2)

    add_body_para(doc,
        '研发费用加计扣除属于企业所得税税基式优惠政策，其基本机制是在企业研发费用据实扣除基础上，再按照'
        '一定比例增加税前扣除额，从而降低企业应纳税所得额和实际税负。2018年，财政部、税务总局和科技部发布'
        '财税〔2018〕99号，将符合条件企业研发费用加计扣除比例提高至75%。2021年，财政部和税务总局公告2021'
        '年第13号明确，制造业企业自2021年1月1日起研发费用按100%加计扣除。2022年，科技型中小企业加计扣除'
        '比例提高至100%。2023年，财政部和税务总局公告2023年第7号规定，符合条件企业研发费用统一按100%加计'
        '扣除，并作为制度性安排长期实施。2023年至2027年，集成电路和工业母机企业适用120%加计扣除。'
    )

    add_body_para(doc,
        '这一政策演进呈现出由普惠性扩围到重点行业加力，再到制度化稳定实施的特征。对本文而言，2021年制造业'
        '100%政策具有较强识别价值，因为它在2021至2022年间形成了制造业与非制造业之间的差异化激励。2023年以后，'
        '符合条件企业普遍适用100%加计扣除，制造业相对于非制造业的政策差异消失，因此2023年以后不适合继续纳入'
        '本文的基准DID样本。'
    )

    # Table 9: Policy evolution
    add_table_title(doc, '表 9 研发费用加计扣除政策演进表')
    policy_data = [
        ['2018', '财税〔2018〕99号', '符合条件企业', '75%', '普惠型研发税收激励增强'],
        ['2021', '财政部 税务总局公告2021年第13号', '制造业企业', '100%', '形成制造业差异化政策冲击'],
        ['2022', '财政部 税务总局 科技部公告2022年第16号', '科技型中小企业', '100%', '支持科技型中小企业研发投入'],
        ['2023', '财政部 税务总局公告2023年第7号', '符合条件企业', '100%', '政策长期化和制度化'],
        ['2023—2027', '财政部等公告2023年第44号', '集成电路和工业母机企业', '120%', '对战略产业叠加定向支持'],
    ]
    table9 = create_three_line_table(doc,
        ['年份', '政策文件', '适用主体', '加计扣除比例', '政策含义'],
        policy_data
    )

    # Insert Figure 6: Policy timeline
    timeline_png = os.path.join(FIGURE_DIR, 'focus_policy_timeline.png')
    if os.path.exists(timeline_png):
        add_figure_title(doc, '图 6 研发费用加计扣除政策时间轴')
        doc.add_picture(timeline_png, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_note(doc, '注：2018至2020年75%；2021至2022年制造业100%；2023年起全行业100%并长期实施。')

    # (二) Theory and mechanism
    add_heading_custom(doc, '（二）理论基础与作用机制', level=2)

    add_body_para(doc,
        '公共支出理论认为，科技创新具有正外部性、高风险和不确定性，企业研发投入的私人收益通常低于社会收益。'
        '在基础研究和共性技术开发领域，知识溢出较强，私人资本容易出现投入不足。政府科技投入的重要作用在于'
        '弥补市场失灵、降低企业研发风险、构建创新生态和推动战略技术突破。'
    )

    add_body_para(doc,
        '研发费用加计扣除作为税收激励工具，主要通过降低企业研发活动的税后边际成本影响研发决策。与直接补贴'
        '相比，税收优惠不直接选择具体研发项目，企业可以根据自身技术路线、产品规划和市场需求决定研发方向，'
        '因而更接近市场友好型激励。税收优惠与直接补贴具有政策组合效应：税收优惠适合作为覆盖面较广、执行成本'
        '较低、对企业自主权干扰较小的效率激励工具，直接补贴则更适合基础共性技术、前沿高风险探索等企业自发'
        '投资意愿较弱的领域。'
    )

    add_body_para(doc,
        '从本文结果看，研发费用加计扣除的政策效果并未主要表现为短期专利数量扩张，而是表现为研发投入相对调整'
        '和单位投入产出效率改善。这与税收优惠工具的属性具有一致性。税收优惠通过降低研发成本和改善企业预期，'
        '使企业能够在较低税后成本下重新配置研发资源，但它并不必然立即增加专利申请数量。专利数量还受到技术'
        '周期、审查周期、企业专利策略和行业研发特征影响，因此政策评估需要引入研发投入结构和创新效率指标。'
    )

    # (三) Macro background
    add_heading_custom(doc, '（三）现实背景与宏观影响', level=2)

    add_body_para(doc,
        '从宏观背景看，我国R&D投入持续增长。国家统计局发布的《2024年全国科技经费投入统计公报》显示，2024年'
        '全国研究与试验发展经费为36326.8亿元，比上年增长8.9%，R&D经费投入强度为2.69%。按执行主体看，企业'
        'R&D经费为28211.6亿元，占全社会R&D经费的77.7%，说明企业仍是我国研发活动的主要承担者。'
    )

    add_body_para(doc,
        '研发费用加计扣除政策覆盖范围也在扩大。据国家税务总局统计，2023年度全国共有62.9万户企业享受研发费用'
        '加计扣除金额3.45万亿元，该政策通过减少应纳税所得额降低企业税负，进而支持企业将节省资金继续投入研发。'
        '宏观事实表明，企业研发投入主体地位不断强化，税收优惠已成为支持企业创新的重要政策工具。本文的微观证据'
        '进一步补充了这一判断：在制造业差异化加计扣除政策背景下，短期内更明显的变化并非专利数量扩张，而是研发'
        '投入相对调整和单位投入产出效率改善。'
    )

    # (四) Policy recommendations
    add_heading_custom(doc, '（四）政策建议', level=2)

    add_body_para(doc,
        '第一，研发费用加计扣除政策应保持稳定性和可预期性。企业研发活动具有较长周期，政策稳定性会影响企业长期'
        '研发规划。2023年将符合条件企业研发费用加计扣除比例统一提高至100%，并作为制度性安排长期实施，有助于'
        '稳定企业研发预期。对于集成电路、工业母机等战略产业，120%加计扣除体现了在普惠政策基础上的定向加力。'
    )

    add_body_para(doc,
        '第二，政策评价体系需要从数量导向转向数量、质量和效率并重。专利数量是衡量创新产出的重要指标，但容易'
        '受到企业申请策略、专利质量和审查周期影响。本文结果显示，制造业专利数量相对变化不显著，但单位研发投入'
        '和单位研发人员创新效率表现出相对改善。因此，后续政策评估应同时关注发明专利质量、被引用次数、研发投入'
        '结构、新产品收入和全要素生产率等指标。'
    )

    add_body_para(doc,
        '第三，税收优惠应与直接补贴、政府采购和产业基金协同配置。对商业化路径较清晰、企业具备自主研发能力的'
        '试验发展活动，加计扣除可以通过降低税负发挥普惠型激励作用。对基础研究、共性关键技术和前沿高风险领域，'
        '直接补贴和政府采购仍具有重要作用。财政科技政策需要根据研发阶段和技术外溢程度配置工具。'
    )

    add_body_para(doc,
        '第四，研发费用归集监管和知识产权保护需要同步强化。税收优惠能否形成有效创新激励，取决于企业研发费用'
        '归集是否真实规范，也取决于创新收益能否得到保护。应完善研发费用辅助账和留存备查制度，加强研发活动'
        '真实性核查。同时，强化知识产权保护能够提高企业创新收益的可占有性，使税收激励更可能转化为长期研发能力。'
    )

    # ============================================================
    # 七、结论
    # ============================================================
    add_heading_custom(doc, '七、结论', level=1)

    add_body_para(doc,
        '本文基于2017至2022年A股上市公司面板数据，以2021年制造业企业研发费用加计扣除比例由75%提高至100%'
        '为政策冲击，采用双重固定效应DID模型考察税收激励政策是否伴随企业研发投入调整和相对创新效率改善。'
        '研究发现，制造业企业研发投入和专利产出在政策后绝对水平上仍保持增长，但相对于非制造业企业，专利数量'
        '变化不显著，研发支出、研发强度和研发人员投入增长显著较慢。在此基础上，单位研发投入和单位研发人员'
        '创新产出效率表现出显著相对改善。效率来源拆解表明，效率改善主要来自研发投入相对增长较慢，而非创新'
        '数量显著增加。'
    )

    add_body_para(doc,
        '异质性分析显示，政策前研发基础较强的企业和制造业内部高研发暴露企业效率改善更明显，说明企业吸收政策'
        '激励并转化为资源配置效率的能力存在差异。稳健性检验表明，加入省份与年份交互固定效应后核心结果保持稳定；'
        '但事件研究和安慰剂检验也显示，研发投入和效率指标在政策前已经存在差异收敛趋势，因而本文不将结果解释为'
        '严格因果意义上的政策效率提升。'
    )

    add_body_para(doc,
        '总体而言，研发费用加计扣除政策的短期影响更适合表述为伴随研发投入调整和相对创新效率改善。该结论提示，'
        '税收激励政策评估需要超越专利数量扩张逻辑，将研发资源配置、单位投入产出效率和政策识别边界纳入分析框架。'
    )

    # ============================================================
    # 参考文献
    # ============================================================
    add_heading_custom(doc, '参考文献', level=1)

    refs = [
        '[1] 财政部，税务总局，科技部. 关于提高研究开发费用税前加计扣除比例的通知：财税〔2018〕99号[EB/OL]. 2018-09-20.',
        '[2] 财政部，税务总局. 关于进一步完善研发费用税前加计扣除政策的公告：财政部 税务总局公告2021年第13号[EB/OL]. 2021-03-31.',
        '[3] 财政部，税务总局，科技部. 关于进一步提高科技型中小企业研发费用税前加计扣除比例的公告：财政部 税务总局 科技部公告2022年第16号[EB/OL]. 2022-03-23.',
        '[4] 财政部，税务总局. 关于进一步完善研发费用税前加计扣除政策的公告：财政部 税务总局公告2023年第7号[EB/OL]. 2023-03-26.',
        '[5] 财政部，税务总局，国家发展改革委，工业和信息化部. 关于提高集成电路和工业母机企业研发费用加计扣除比例的公告：2023年第44号[EB/OL]. 2023-09-12.',
        '[6] 国家统计局，科学技术部，财政部. 2024年全国科技经费投入统计公报[EB/OL]. 2025-09-29.',
        '[7] 新华社. 关于"研发费用加计扣除"，你了解吗？[EB/OL]. 2024-10-02.',
        '[8] 王文. 专题分析：科技支出[Z]. 中央财经大学公共支出分析课程课件.',
    ]
    for ref in refs:
        add_paragraph_with_font(doc, ref, cn_font='宋体', en_font='Times New Roman',
                                size=Pt(10.5), line_spacing=1.25, space_after=Pt(2))

    # ============================================================
    # 附录
    # ============================================================
    doc.add_page_break()
    add_heading_custom(doc, '附录', level=1)

    # 附录A
    add_heading_custom(doc, '附录A 数据与实验来源考证', level=2)
    add_body_para(doc,
        '本文实验结果来自本地项目/home/u2023312303/裴的实验/tec_expenditure。企业面板数据文件包括'
        'firm_panel_v3.csv和firm_panel_v4.csv，回归结果文件包括focus_quantity_results.csv、'
        'focus_rd_adjustment_results.csv、focus_efficiency_main_results.csv、focus_efficiency_'
        'decomposition.csv、focus_high_rd_heterogeneity.csv、focus_within_manufacturing_exposure.csv、'
        'focus_robustness_summary.csv、focus_event_study.csv和focus_placebo.csv。主分析脚本为'
        'run_focus_analysis.py，使用Python 3.12、linearmodels和statsmodels，可通过"uv run python '
        'run_focus_analysis.py"重新运行。所有核心输出文件均存在，图表文件位于outputs/figures/目录。'
    )

    # 附录B
    add_heading_custom(doc, '附录B 政策来源考证', level=2)
    add_body_para(doc,
        '研发费用加计扣除政策来源包括：财税〔2018〕99号将符合条件企业研发费用加计扣除比例提高至75%；'
        '财政部税务总局公告2021年第13号将制造业企业加计扣除比例提高至100%；财政部税务总局科技部公告2022年'
        '第16号将科技型中小企业加计扣除比例提高至100%；财政部税务总局公告2023年第7号将符合条件企业加计扣除'
        '比例统一提高至100%；财政部等公告2023年第44号将集成电路和工业母机企业加计扣除比例提高至120%（2023'
        '至2027年）。政策文件均可在国家税务总局政策法规库（https://fgk.chinatax.gov.cn）查阅。'
    )

    # 附录C
    add_heading_custom(doc, '附录C 课件理论来源考证', level=2)
    add_body_para(doc,
        '课程课件《专题分析：科技支出》提供了本文政策分析的理论依据。课件指出，科技创新具有正外部性、高风险'
        '和公共品属性，政府科技投入可以弥补市场失灵、降低企业研发风险并推动战略技术突破。课件还强调，税收优惠'
        '与直接补贴具有政策组合效应，税收优惠更适合覆盖面广、执行成本较低、企业自主权较强的创新活动，直接补贴'
        '更适合基础共性技术和高风险前沿领域。'
    )

    # 附录D: Supplementary tables
    add_heading_custom(doc, '附录D 补充统计表', level=2)

    # 附表A: 完整变量定义表
    add_table_title(doc, '附表A 完整变量定义表')
    appendix_vars = [
        ['被解释变量', '发明专利申请', 'ln_invention_apply', 'ln(1+发明专利申请数量)', 'CSMAR专利表', 'log1p变换，1%/99%缩尾'],
        ['被解释变量', '发明专利授权', 'ln_invention_grant', 'ln(1+发明专利授权数量)', 'CSMAR专利表', 'log1p变换，1%/99%缩尾'],
        ['被解释变量', '专利总申请', 'ln_patent_apply_total', 'ln(1+三类专利合计)', 'CSMAR专利表', 'log1p变换，1%/99%缩尾'],
        ['被解释变量', '专利总授权', 'ln_patent_grant_total', 'ln(1+三类授权合计)', 'CSMAR专利表', 'log1p变换，1%/99%缩尾'],
        ['被解释变量', '研发支出', 'ln_rd_expense_10k', 'ln(1+研发支出/10000)', 'CSMAR研发投入表', '元转万元+log1p'],
        ['被解释变量', '研发强度', 'rd_intensity_01', '研发支出/营业收入(0—1)', 'CSMAR研发投入表', '百分比÷100'],
        ['被解释变量', '研发人员', 'ln_rd_staff', 'ln(1+研发人员数量)', 'CSMAR研发投入表', 'log1p变换'],
        ['被解释变量', '研发人员占比', 'rd_staff_ratio_01', '研发人员/员工总数(0—1)', 'CSMAR研发投入表', '百分比÷100'],
        ['核心解释变量', '制造业×政策后', 'manufacturing_post2021', 'manufacturing × post2021', '构造', '交互项'],
        ['控制变量', '企业规模', 'ln_assets', 'ln总资产', 'CSMAR资产负债表', '对数变换'],
        ['控制变量', '盈利能力', 'roa', '净利润/总资产', 'CSMAR利润表+资产负债表', '比率，1%/99%缩尾'],
        ['控制变量', '现金流比率', 'cashflow_ratio', '经营现金流/总资产', 'CSMAR现金流量表', '比率，1%/99%缩尾'],
        ['控制变量', '企业年龄', 'firm_age', '年份−成立年份', 'CSMAR基本信息表', '1%/99%缩尾'],
    ]
    table_a1 = create_three_line_table(doc,
        ['变量类型', '变量名称', '变量符号', '变量定义', '数据来源', '处理方式'],
        appendix_vars
    )

    # 附表B: Manufacturing vs non-manufacturing
    add_table_title(doc, '附表B 制造业与非制造业分组描述性统计表')
    group_data = [
        ['ln_invention_apply', '0.512', '0.181', '0.331', '27.547', '0.0000***'],
        ['ln_invention_grant', '0.381', '0.161', '0.220', '22.531', '0.0000***'],
        ['ln_patent_apply_total', '0.762', '0.277', '0.485', '30.082', '0.0000***'],
        ['ln_rd_expense_10k', '8.672', '5.545', '3.128', '68.850', '0.0000***'],
        ['rd_intensity_01', '0.054', '0.055', '−0.000', '−0.380', '0.7036'],
        ['ln_rd_staff', '5.181', '3.260', '1.922', '61.626', '0.0000***'],
        ['eff_apply_rd_10k', '−8.160', '−5.363', '−2.797', '−62.112', '0.0000***'],
        ['ln_assets', '21.965', '22.638', '−0.673', '−32.435', '0.0000***'],
        ['roa', '0.049', '0.028', '0.021', '18.886', '0.0000***'],
        ['cashflow_ratio', '0.053', '0.041', '0.012', '13.265', '0.0000***'],
    ]
    table_a2 = create_three_line_table(doc,
        ['变量', '制造业均值', '非制造业均值', '均值差异', 't值', 'p值'],
        group_data
    )
    add_table_note(doc, '注：*** p<0.01。t检验采用Welch近似。')

    # 附表C: Ownership heterogeneity
    add_table_title(doc, '附表C 所有制异质性完整结果表')
    soe_data = [
        ['eff_apply_rd_10k', '交互项−非国企', 'did_x_private', '0.167', '0.1000', '0.0952*', '25,159'],
        ['eff_grant_rd_10k', '交互项−非国企', 'did_x_private', '0.007', '0.0980', '0.9415', '25,159'],
        ['eff_apply_staff', '交互项−非国企', 'did_x_private', '0.076', '0.1095', '0.4899', '25,159'],
        ['eff_grant_staff', '交互项−非国企', 'did_x_private', '−0.084', '0.1052', '0.4241', '25,159'],
        ['eff_apply_rd_10k', '分组−国企', 'manufacturing_post2021', '0.269', '0.0658', '0.0000***', '23,028'],
        ['eff_grant_rd_10k', '分组−国企', 'manufacturing_post2021', '0.272', '0.0656', '0.0000***', '23,028'],
        ['eff_apply_rd_10k', '分组−非国企', 'manufacturing_post2021', '0.210', '0.2231', '0.3456', '2,131'],
        ['eff_grant_staff', '分组−非国企', 'manufacturing_post2021', '0.434', '0.1842', '0.0186**', '2,131'],
    ]
    table_a3 = create_three_line_table(doc,
        ['因变量', '模型', '系数项', '系数', '标准误', 'p值', 'N'],
        soe_data
    )
    add_table_note(doc, '注：非国企分组样本较小（524 firms, 2,131 obs），估计精度较低。*** p<0.01，** p<0.05，* p<0.10。')

    # 附表D: Placebo full results
    add_table_title(doc, '附表D 安慰剂检验完整结果表')
    placebo_data = [
        ['eff_apply_rd_10k', '2019', '0.290***', '0.0719', '0.0001', '15,837'],
        ['eff_apply_rd_10k', '2020', '0.325***', '0.0701', '0.0000', '15,837'],
        ['eff_grant_rd_10k', '2019', '0.237***', '0.0715', '0.0009', '15,837'],
        ['eff_grant_rd_10k', '2020', '0.296***', '0.0688', '0.0000', '15,837'],
        ['eff_apply_staff', '2019', '0.116**', '0.0489', '0.0180', '15,837'],
        ['eff_apply_staff', '2020', '0.157***', '0.0528', '0.0029', '15,837'],
        ['eff_grant_staff', '2019', '0.062', '0.0483', '0.1995', '15,837'],
        ['eff_grant_staff', '2020', '0.128**', '0.0511', '0.0125', '15,837'],
        ['ln_invention_apply', '2019', '0.024', '0.0262', '0.3597', '15,837'],
        ['ln_invention_apply', '2020', '0.029', '0.0291', '0.3260', '15,837'],
        ['ln_rd_expense_10k', '2019', '−0.266***', '0.0676', '0.0001', '15,837'],
        ['ln_rd_expense_10k', '2020', '−0.297***', '0.0642', '0.0000', '15,837'],
    ]
    table_a4 = create_three_line_table(doc,
        ['因变量', '假想政策时点', '系数', '标准误', 'p值', 'N'],
        placebo_data
    )
    add_table_note(doc,
        '注：安慰剂检验使用2017至2020年子样本。假想2019年时点使用2017至2020年样本，2019至2020年定义为"政策后"。'
        '假想2020年时点使用2017至2020年样本，2020年定义为"政策后"。*** p<0.01，** p<0.05。'
    )

    # 附表E: Robustness full summary
    add_table_title(doc, '附表E 稳健性检验完整结果表')
    robust_data = [
        ['Prov×Year FE', 'eff_apply_rd_10k', '0.264***', '0.255***', '0.0635', '0.0001', '保持显著'],
        ['Prov×Year FE', 'eff_grant_rd_10k', '0.277***', '0.290***', '0.0632', '0.0000', '保持显著'],
        ['Prov×Year FE', 'eff_apply_staff', '0.172***', '0.161***', '0.0512', '0.0016', '保持显著'],
        ['Prov×Year FE', 'eff_grant_staff', '0.186***', '0.196***', '0.0506', '0.0001', '保持显著'],
        ['Prov×Year FE', 'ln_rd_expense_10k', '−0.288***', '−0.291***', '0.0587', '0.0000', '保持显著'],
        ['Prov×Year FE', 'ln_rd_staff', '−0.197***', '−0.197***', '0.0452', '0.0000', '保持显著'],
        ['asinh比率', 'asinh_apply_per_rd', '—', '−8.52×10⁻⁹', '3.94×10⁻⁹', '0.0305', '方向相反'],
        ['asinh比率', 'asinh_grant_per_rd', '—', '−1.79×10⁻⁹', '1.81×10⁻⁹', '0.3235', '不显著'],
        ['asinh比率', 'asinh_grant_per_staff', '—', '9.59×10⁻⁵', '0.0006', '0.8729', '不显著'],
    ]
    table_a5 = create_three_line_table(doc,
        ['检验类型', '因变量', '基准系数', '稳健性系数', '标准误', 'p值', '结论'],
        robust_data
    )
    add_table_note(doc, '注：asinh比率型效率指标结果不理想，正文中仅作为稳健性讨论。*** p<0.01。')

    # ============================================================
    # Save
    # ============================================================
    output_path = os.path.join(OUTPUT_DIR, 'paper_innovation_efficiency.docx')
    doc.save(output_path)
    print(f'论文Word文档已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_paper()
